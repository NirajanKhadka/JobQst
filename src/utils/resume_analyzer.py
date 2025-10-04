#!/usr/bin/env python3
"""
Automated Resume Analyzer for JobQst.
This module automatically extracts keywords, skills, and experience level from resumes
to enable Configurable, targeted job scraping.
"""

import re
import json
from pathlib import Path
from typing import Dict, List, Set, Tuple
from rich.console import Console
import docx
from docx import Document
from src.utils.profile_helpers import load_profile

console = Console()


class ResumeAnalyzer:
    """
    Automated resume analyzer that extracts relevant keywords and skills
    for targeted job searching.
    """

    def __init__(self, profile_name: str = "default"):
        self.profile_name = profile_name
        # Technical skills database
        self.technical_skills = {
            # Programming Languages
            "programming": [
                "python",
                "java",
                "javascript",
                "typescript",
                "c++",
                "c#",
                "php",
                "ruby",
                "go",
                "rust",
                "swift",
                "kotlin",
                "scala",
                "r",
                "matlab",
                "sas",
                "vba",
                "html",
                "css",
                "sql",
            ],
            # Data & Analytics
            "data_analytics": [
                "pandas",
                "numpy",
                "scipy",
                "matplotlib",
                "seaborn",
                "plotly",
                "tableau",
                "power bi",
                "excel",
                "google analytics",
                "statistical analysis",
                "data visualization",
                "data mining",
                "machine learning",
                "deep learning",
                "artificial intelligence",
                "predictive modeling",
                "regression analysis",
                "time series",
                "a/b testing",
                "hypothesis testing",
            ],
            # Databases
            "databases": [
                "mysql",
                "postgresql",
                "mongodb",
                "redis",
                "elasticsearch",
                "oracle",
                "sql server",
                "sqlite",
                "cassandra",
                "dynamodb",
                "snowflake",
                "bigquery",
            ],
            # Cloud & DevOps
            "cloud_devops": [
                "aws",
                "azure",
                "google cloud",
                "gcp",
                "docker",
                "kubernetes",
                "jenkins",
                "git",
                "github",
                "gitlab",
                "terraform",
                "ansible",
                "linux",
                "bash",
                "powershell",
            ],
            # Web Development
            "web_development": [
                "react",
                "angular",
                "vue",
                "node.js",
                "express",
                "django",
                "flask",
                "spring",
                "bootstrap",
                "jquery",
                "rest api",
                "graphql",
                "microservices",
            ],
            # Business & Soft Skills
            "business": [
                "project management",
                "agile",
                "scrum",
                "business analysis",
                "requirements gathering",
                "stakeholder management",
                "process improvement",
                "quality assurance",
                "testing",
                "documentation",
                "communication",
                "leadership",
                "teamwork",
            ],
        }

        # Job title patterns for experience level detection
        self.experience_indicators = {
            "entry": [
                "junior",
                "entry",
                "associate",
                "trainee",
                "intern",
                "graduate",
                "assistant",
                "coordinator",
                "analyst i",
                "developer i",
                "level 1",
                "tier 1",
            ],
            "mid": ["analyst ii", "developer ii", "specialist", "consultant", "level 2", "tier 2"],
            "senior": [
                "senior",
                "lead",
                "principal",
                "manager",
                "director",
                "head",
                "chief",
                "architect",
                "expert",
                "level 3",
                "tier 3",
            ],
        }

        # Industry-specific job titles
        self.job_titles = [
            "data analyst",
            "business analyst",
            "financial analyst",
            "research analyst",
            "data scientist",
            "data engineer",
            "business intelligence analyst",
            "software developer",
            "software engineer",
            "web developer",
            "full stack developer",
            "python developer",
            "java developer",
            "frontend developer",
            "backend developer",
            "machine learning engineer",
            "ai engineer",
            "devops engineer",
            "cloud engineer",
            "project manager",
            "product manager",
            "scrum master",
            "business consultant",
            "quality assurance",
            "qa engineer",
            "test engineer",
            "systems analyst",
        ]

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text content from a DOCX file."""
        try:
            doc = Document(docx_path)
            text = []

            for paragraph in doc.paragraphs:
                text.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)

            return "\n".join(text)

        except Exception as e:
            console.print(f"[red]❌ Error reading DOCX file: {e}[/red]")
            return ""

    def extract_skills(self, resume_text: str) -> Dict[str, List[str]]:
        """Extract technical skills from resume text."""
        resume_lower = resume_text.lower()
        found_skills = {}

        for category, skills in self.technical_skills.items():
            found_skills[category] = []

            for skill in skills:
                # Use word boundaries to avoid partial matches
                pattern = r"\b" + re.escape(skill.lower()) + r"\b"
                if re.search(pattern, resume_lower):
                    found_skills[category].append(skill.title())

        # Remove empty categories
        found_skills = {k: v for k, v in found_skills.items() if v}

        return found_skills

    def extract_job_titles(self, resume_text: str) -> List[str]:
        """Extract relevant job titles from resume text."""
        resume_lower = resume_text.lower()
        found_titles = []

        for title in self.job_titles:
            pattern = r"\b" + re.escape(title.lower()) + r"\b"
            if re.search(pattern, resume_lower):
                found_titles.append(title.title())

        return found_titles

    def determine_experience_level(self, resume_text: str) -> str:
        """Determine experience level from resume content."""
        resume_lower = resume_text.lower()

        # Look for years of experience
        year_patterns = [
            r"(\d+)\+?\s*years?\s*(?:of\s*)?experience",
            r"(\d+)\+?\s*years?\s*in",
            r"over\s*(\d+)\s*years?",
            r"more\s*than\s*(\d+)\s*years?",
        ]

        max_years = 0
        for pattern in year_patterns:
            matches = re.findall(pattern, resume_lower)
            for match in matches:
                years = int(match)
                max_years = max(max_years, years)

        # Classify based on years
        if max_years == 0:
            # Look for experience level indicators in text
            for level, indicators in self.experience_indicators.items():
                for indicator in indicators:
                    if indicator in resume_lower:
                        return level
            return "entry"  # Default for new graduates
        elif max_years <= 2:
            return "entry"
        elif max_years <= 5:
            return "mid"
        else:
            return "senior"

    def generate_search_keywords(
        self, skills: Dict[str, List[str]], job_titles: List[str], experience_level: str
    ) -> List[str]:
        """Generate Automated search keywords based on extracted information."""
        keywords = []

        # Add experience-appropriate job titles
        if experience_level == "entry":
            prefixes = ["Junior", "Entry Level", "Associate", "Graduate"]
            for title in job_titles[:3]:  # Top 3 relevant titles
                keywords.append(title)
                for prefix in prefixes:
                    keywords.append(f"{prefix} {title}")
        else:
            keywords.extend(job_titles[:5])

        # Add top technical skills as search terms
        all_skills = []
        for skill_list in skills.values():
            all_skills.extend(skill_list)

        # Prioritize programming languages and key tools
        priority_skills = []
        if "programming" in skills:
            priority_skills.extend(skills["programming"][:3])
        if "data_analytics" in skills:
            priority_skills.extend(skills["data_analytics"][:3])

        keywords.extend(priority_skills)

        # Add combination keywords for better targeting
        if "Python" in all_skills and "Data Analyst" in [t.title() for t in job_titles]:
            keywords.append("Python Data Analyst")
        if "SQL" in all_skills:
            keywords.append("SQL Analyst")

        # Remove duplicates and limit
        keywords = list(dict.fromkeys(keywords))  # Preserve order while removing duplicates
        return keywords[:10]  # Limit to top 10 keywords

    def analyze_resume(self, profile: Dict) -> Dict:
        """
        Analyze resume and generate Automated keywords and job targeting info.
        """
        console.print("[bold blue]🧠 Analyzing resume for Automated job targeting...[/bold blue]")

        # Get resume path
        profile_dir = Path(profile.get("profile_dir", f"profiles/{profile['profile_name']}"))
        resume_docx = profile_dir / profile.get(
            "resume_docx", f"{profile['profile_name']}_Resume.docx"
        )

        if not resume_docx.exists():
            console.print(f"[red]❌ Resume not found: {resume_docx}[/red]")
            return self._fallback_analysis(profile)

        # Extract text from resume
        console.print(f"[cyan]📄 Reading resume: {resume_docx}[/cyan]")
        resume_text = self.extract_text_from_docx(str(resume_docx))

        if not resume_text:
            console.print("[red]❌ Could not extract text from resume[/red]")
            return self._fallback_analysis(profile)

        console.print(f"[green]✅ Extracted {len(resume_text)} characters from resume[/green]")

        # Analyze resume content
        skills = self.extract_skills(resume_text)
        job_titles = self.extract_job_titles(resume_text)
        experience_level = self.determine_experience_level(resume_text)

        # Generate Automated keywords
        keywords = self.generate_search_keywords(skills, job_titles, experience_level)

        # Create analysis result
        analysis = {
            "extracted_skills": skills,
            "relevant_job_titles": job_titles,
            "experience_level": experience_level,
            "Automated_keywords": keywords,
            "resume_length": len(resume_text),
            "analysis_timestamp": Path(__file__).stat().st_mtime,
        }

        # Display results
        self._display_analysis_results(analysis)

        return analysis

    def _fallback_analysis(self, profile: Dict) -> Dict:
        """Fallback analysis using existing profile data."""
        console.print("[yellow]⚠️ Using fallback analysis from profile data[/yellow]")

        existing_keywords = profile.get("keywords", [])
        existing_skills = profile.get("skills", [])

        return {
            "extracted_skills": {"existing": existing_skills},
            "relevant_job_titles": existing_keywords,
            "experience_level": "entry",  # Default assumption
            "Automated_keywords": existing_keywords[:10],
            "resume_length": 0,
            "analysis_timestamp": 0,
        }

    def _display_analysis_results(self, analysis: Dict):
        """Display analysis results in a nice format."""
        console.print("\n[bold green]📊 Resume Analysis Results:[/bold green]")

        # Experience level
        exp_level = analysis["experience_level"]
        exp_color = "green" if exp_level == "entry" else "yellow" if exp_level == "mid" else "red"
        console.print(
            f"[bold]🎯 Experience Level:[/bold] [{exp_color}]{exp_level.upper()}[/{exp_color}]"
        )

        # Skills by category
        skills = analysis["extracted_skills"]
        if skills:
            console.print(f"\n[bold]🛠️ Technical Skills Found:[/bold]")
            for category, skill_list in skills.items():
                category_name = category.replace("_", " ").title()
                console.print(f"   [cyan]{category_name}:[/cyan] {', '.join(skill_list[:5])}")
                if len(skill_list) > 5:
                    console.print(f"      ... and {len(skill_list) - 5} more")

        # Job titles
        job_titles = analysis["relevant_job_titles"]
        if job_titles:
            console.print(f"\n[bold]💼 Relevant Job Titles:[/bold]")
            for title in job_titles[:5]:
                console.print(f"   • {title}")

        # Generated keywords
        keywords = analysis["Automated_keywords"]
        console.print(f"\n[bold]🔍 Generated Search Keywords:[/bold]")
        for i, keyword in enumerate(keywords, 1):
            console.print(f"   {i}. [green]{keyword}[/green]")

        console.print(f"\n[cyan]📈 Total keywords generated: {len(keywords)}[/cyan]")


def analyze_profile_resume(profile: Dict) -> Dict:
    """Convenience function to analyze a profile's resume."""
    analyzer = ResumeAnalyzer()
    return analyzer.analyze_resume(profile)

    # Main functionality moved to CLI module or tests
    # Import and use the functions directly

    try:
        profile = profile_helpers.load_profile("Nirajan")
        analyzer = ResumeAnalyzer()
        analysis = analyzer.analyze_resume(profile)

        console.print(f"\n[bold]🎉 Analysis complete![/bold]")
        console.print(f"Generated {len(analysis['Automated_keywords'])} Automated keywords")

    except Exception as e:
        console.print(f"[red]❌ Error: {e}[/red]")
        import traceback

        traceback.print_exc()
