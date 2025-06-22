import json
import os
import re
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from rich.console import Console

from src.core import utils

# Initialize console for rich output
console = Console()

# Try to import ollama and check if it's working
OLLAMA_AVAILABLE = False
OLLAMA_MODULE = None

def fix_ssl_cert_path():
    """Fix SSL certificate path for Ollama import."""
    import os
    import sys

    # Check if SSL_CERT_FILE is set incorrectly
    ssl_cert_file = os.environ.get('SSL_CERT_FILE', '')

    if ssl_cert_file and not ssl_cert_file.endswith('.pem'):
        # If it points to a directory, try to find the cert file
        potential_cert_path = os.path.join(ssl_cert_file, 'cacert.pem')
        if os.path.exists(potential_cert_path):
            os.environ['SSL_CERT_FILE'] = potential_cert_path
            console.print(f"[cyan]🔧 Fixed SSL_CERT_FILE path: {potential_cert_path}[/cyan]")
            return True
        else:
            # Try common conda cert locations
            if 'miniconda3' in ssl_cert_file or 'anaconda3' in ssl_cert_file:
                # Extract conda env path and try standard locations
                env_path = ssl_cert_file.replace('\\Library\\ssl', '').replace('/Library/ssl', '')
                potential_paths = [
                    os.path.join(env_path, 'Library', 'ssl', 'cacert.pem'),
                    os.path.join(env_path, 'Library', 'ssl', 'cert.pem'),
                    os.path.join(env_path, 'ssl', 'cacert.pem'),
                ]

                for path in potential_paths:
                    if os.path.exists(path):
                        os.environ['SSL_CERT_FILE'] = path
                        console.print(f"[cyan]🔧 Fixed SSL_CERT_FILE path: {path}[/cyan]")
                        return True

    return False

try:
    # First, try to fix SSL certificate path
    fix_ssl_cert_path()

    import ollama
    OLLAMA_MODULE = ollama

    # Test if Ollama is actually working by trying to list models
    try:
        # Use a more robust way to test Ollama availability
        import requests
        response = requests.get("http://localhost:11434/api/tags", timeout=3)
        if response.status_code == 200:
            models_data = response.json()
            if models_data and 'models' in models_data and models_data['models']:
                OLLAMA_AVAILABLE = True
                console.print("[green]✅ Ollama is available and working[/green]")
            else:
                console.print("[yellow]⚠️ Ollama running but no models found[/yellow]")
                console.print("[yellow]💡 Try: ollama pull mistral[/yellow]")
        else:
            console.print("[yellow]⚠️ Ollama service not responding[/yellow]")
            console.print("[yellow]💡 Try running: ollama serve[/yellow]")
    except requests.exceptions.RequestException:
        console.print("[yellow]⚠️ Ollama service not running[/yellow]")
        console.print("[yellow]💡 Try running: ollama serve[/yellow]")
    except Exception as e:
        console.print(f"[yellow]⚠️ Error checking Ollama status: {e}[/yellow]")

except ImportError:
    console.print("[yellow]⚠️ Ollama Python package not installed[/yellow]")
    console.print("[yellow]💡 Install with: pip install ollama[/yellow]")
except OSError as e:
    if "Invalid argument" in str(e) and "SSL_CERT_FILE" in str(e):
        console.print(f"[yellow]⚠️ Ollama import error: {e}[/yellow]")
        console.print("[yellow]💡 SSL certificate path issue detected[/yellow]")
        console.print("[yellow]💡 Try running: unset SSL_CERT_FILE[/yellow]")
        console.print("[yellow]💡 Or fix with: export SSL_CERT_FILE=/path/to/cacert.pem[/yellow]")
    else:
        console.print(f"[yellow]⚠️ Ollama import error: {e}[/yellow]")
        console.print("[yellow]💡 This may be due to SSL certificate issues[/yellow]")
except Exception as e:
    console.print(f"[yellow]⚠️ Ollama import error: {e}[/yellow]")
    console.print("[yellow]💡 This may be due to SSL certificate issues[/yellow]")

if not OLLAMA_AVAILABLE:
    console.print("[cyan]📝 Document customization will use fallback methods[/cyan]")

def customize(job: Dict, profile: Dict) -> Tuple[str, str]:
    """
    Customize resume and cover letter for a specific job with comprehensive fallback methods.
    Fallback chain: AI Enhancement -> Template Customization -> Basic Replacement -> Emergency Text

    Args:
        job: Job dictionary with title, company, summary, and keywords
        profile: User profile dictionary

    Returns:
        Tuple of (resume_path, cover_letter_path)
    """
    console.print(f"[bold]Customizing documents for:[/bold] {job['title']} at {job['company']}")

    # Get job hash for document naming
    job_hash = utils.hash_job(job)

    # Try to customize resume with fallback methods
    resume_path = customize_resume_with_fallbacks(job, profile, job_hash)

    # Try to customize cover letter with fallback methods
    cover_letter_path = customize_cover_letter_with_fallbacks(job, profile, job_hash)

    return resume_path, cover_letter_path


def customize_resume_with_fallbacks(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Customize resume with comprehensive fallback methods.

    Args:
        job: Job dictionary
        profile: User profile dictionary
        job_hash: Hash of the job for document naming

    Returns:
        Path to the customized resume
    """
    console.print("[green]Customizing resume with fallback methods...[/green]")

    # Method 1: Try AI-enhanced customization
    try:
        return customize_resume_ai_enhanced(job, profile, job_hash)
    except Exception as e:
        console.print(f"[yellow]AI-enhanced resume failed: {e}[/yellow]")

    # Method 2: Try template-based customization
    try:
        return customize_resume_template_based(job, profile, job_hash)
    except Exception as e:
        console.print(f"[yellow]Template-based resume failed: {e}[/yellow]")

    # Method 3: Try basic keyword replacement
    try:
        return customize_resume_basic_replacement(job, profile, job_hash)
    except Exception as e:
        console.print(f"[yellow]Basic replacement resume failed: {e}[/yellow]")

    # Method 4: Emergency fallback - create basic text resume
    try:
        return create_emergency_text_resume(job, profile, job_hash)
    except Exception as e:
        console.print(f"[red]Emergency text resume failed: {e}[/red]")
        raise RuntimeError("All resume customization methods failed")


def customize_cover_letter_with_fallbacks(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Customize cover letter with comprehensive fallback methods.

    Args:
        job: Job dictionary
        profile: User profile dictionary
        job_hash: Hash of the job for document naming

    Returns:
        Path to the customized cover letter
    """
    console.print("[green]Customizing cover letter with fallback methods...[/green]")

    # Method 1: Try AI-enhanced customization
    try:
        return customize_cover_letter_ai_enhanced(job, profile, job_hash)
    except Exception as e:
        console.print(f"[yellow]AI-enhanced cover letter failed: {e}[/yellow]")

    # Method 2: Try template-based customization
    try:
        return customize_cover_letter_template_based(job, profile, job_hash)
    except Exception as e:
        console.print(f"[yellow]Template-based cover letter failed: {e}[/yellow]")

    # Method 3: Try basic placeholder replacement
    try:
        return customize_cover_letter_basic_replacement(job, profile, job_hash)
    except Exception as e:
        console.print(f"[yellow]Basic replacement cover letter failed: {e}[/yellow]")

    # Method 4: Emergency fallback - create basic text cover letter
    try:
        return create_emergency_text_cover_letter(job, profile, job_hash)
    except Exception as e:
        console.print(f"[red]Emergency text cover letter failed: {e}[/red]")
        raise RuntimeError("All cover letter customization methods failed")


def customize_resume_ai_enhanced(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    AI-enhanced resume customization (Method 1).

    Args:
        job: Job dictionary
        profile: User profile dictionary
        job_hash: Hash of the job for document naming

    Returns:
        Path to the customized resume
    """
    if not OLLAMA_AVAILABLE:
        raise RuntimeError("Ollama not available for AI enhancement")

    return customize_resume(job, profile, job_hash)


def customize_resume_template_based(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Template-based resume customization (Method 2).

    Args:
        job: Job dictionary
        profile: User profile dictionary
        job_hash: Hash of the job for document naming

    Returns:
        Path to the customized resume
    """
    console.print("[cyan]Using template-based resume customization[/cyan]")

    # Load template resume
    profile_dir = Path(profile["profile_dir"])
    resume_docx = profile_dir / profile["resume_docx"]

    if not resume_docx.exists():
        raise FileNotFoundError(f"Resume template not found: {resume_docx}")

    # Load the document
    doc = Document(resume_docx)

    # Get missing keywords (simplified approach)
    profile_skills = set(k.lower() for k in profile.get("skills", []))
    job_keywords = set(k.lower() for k in job.get("keywords", []))
    missing_keywords = job_keywords - profile_skills

    if missing_keywords:
        console.print(f"[yellow]Adding missing keywords: {', '.join(list(missing_keywords)[:3])}[/yellow]")
        # Add keywords using template-based approach (no AI)
        add_keywords_template_based(doc, list(missing_keywords)[:3])

    # Save to temp file
    temp_docx = utils.create_temp_file(prefix=f"resume_{job_hash}_", suffix=".docx")
    doc.save(temp_docx)

    # Save customized document and convert to PDF
    output_path = utils.save_customized_document(profile, job_hash, temp_docx, is_resume=True)

    return output_path


def customize_resume_basic_replacement(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Basic keyword replacement resume customization (Method 3).

    Args:
        job: Job dictionary
        profile: User profile dictionary
        job_hash: Hash of the job for document naming

    Returns:
        Path to the customized resume
    """
    console.print("[cyan]Using basic replacement resume customization[/cyan]")

    # Load template resume
    profile_dir = Path(profile["profile_dir"])
    resume_docx = profile_dir / profile["resume_docx"]

    if not resume_docx.exists():
        raise FileNotFoundError(f"Resume template not found: {resume_docx}")

    # Load the document
    doc = Document(resume_docx)

    # Simple placeholder replacement
    replacements = {
        "{{JOB_TITLE}}": job.get("title", ""),
        "{{COMPANY}}": job.get("company", ""),
        "{{LOCATION}}": job.get("location", ""),
        "{{NAME}}": profile.get("name", ""),
        "{{EMAIL}}": profile.get("email", ""),
        "{{PHONE}}": profile.get("phone", "")
    }

    replace_placeholders(doc, replacements)

    # Save to temp file
    temp_docx = utils.create_temp_file(prefix=f"resume_{job_hash}_", suffix=".docx")
    doc.save(temp_docx)

    # Save customized document and convert to PDF
    output_path = utils.save_customized_document(profile, job_hash, temp_docx, is_resume=True)

    return output_path


def create_emergency_text_resume(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Emergency text-based resume creation (Method 4).

    Args:
        job: Job dictionary
        profile: User profile dictionary
        job_hash: Hash of the job for document naming

    Returns:
        Path to the emergency resume
    """
    console.print("[red]Creating emergency text resume[/red]")

    # Create a basic text resume
    resume_content = f"""
{profile.get('name', 'Your Name')}
{profile.get('email', 'your.email@example.com')} | {profile.get('phone', 'Your Phone')}
{profile.get('location', 'Your Location')}

OBJECTIVE
Seeking the position of {job.get('title', 'Job Title')} at {job.get('company', 'Company Name')}.

SKILLS
{', '.join(profile.get('skills', ['Skill 1', 'Skill 2', 'Skill 3']))}
{', '.join(job.get('keywords', [])[:5])}

EXPERIENCE
• Relevant experience in {job.get('title', 'the field')}
• Proficient in {', '.join(job.get('keywords', [])[:3])}
• Strong analytical and problem-solving skills

EDUCATION
• Relevant educational background
• Continuous learning and professional development

This resume was generated automatically for the position at {job.get('company', 'the company')}.
"""

    # Save as text file
    temp_txt = utils.create_temp_file(prefix=f"resume_{job_hash}_", suffix=".txt")
    with open(temp_txt, 'w', encoding='utf-8') as f:
        f.write(resume_content)

    console.print(f"[yellow]Emergency text resume created: {temp_txt}[/yellow]")
    return temp_txt


def add_keywords_template_based(doc: Document, keywords: List[str]) -> None:
    """
    Add keywords using template-based approach (no AI).

    Args:
        doc: Document object
        keywords: List of keywords to add
    """
    if not keywords:
        return

    # Find the last bullet point and add keywords there
    bullet_paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip().startswith('•') or para.text.strip().startswith('-'):
            bullet_paragraphs.append(para)

    if bullet_paragraphs:
        # Add to the last bullet point
        last_bullet = bullet_paragraphs[-1]
        for keyword in keywords[:2]:  # Add max 2 keywords
            if keyword.lower() not in last_bullet.text.lower():
                last_bullet.text = f"{last_bullet.text.rstrip('.')}. Experience with {keyword}."
    else:
        # Add a new skills section
        add_skills_section(doc, set(keywords))


def customize_cover_letter_ai_enhanced(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    AI-enhanced cover letter customization (Method 1).

    Args:
        job: Job dictionary
        profile: User profile dictionary
        job_hash: Hash of the job for document naming

    Returns:
        Path to the customized cover letter
    """
    if not OLLAMA_AVAILABLE:
        raise RuntimeError("Ollama not available for AI enhancement")

    return customize_cover_letter(job, profile, job_hash)


def customize_cover_letter_template_based(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Template-based cover letter customization (Method 2).

    Args:
        job: Job dictionary
        profile: User profile dictionary
        job_hash: Hash of the job for document naming

    Returns:
        Path to the customized cover letter
    """
    console.print("[cyan]Using template-based cover letter customization[/cyan]")

    # Load template cover letter
    profile_dir = Path(profile["profile_dir"])
    cover_letter_docx = profile_dir / profile["cover_letter_docx"]

    if not cover_letter_docx.exists():
        raise FileNotFoundError(f"Cover letter template not found: {cover_letter_docx}")

    # Load the document
    doc = Document(cover_letter_docx)

    # Replace placeholders
    replace_placeholders(doc, {
        "{{JOB_TITLE}}": job.get("title", ""),
        "{{COMPANY}}": job.get("company", ""),
        "{{LOCATION}}": job.get("location", ""),
        "{{NAME}}": profile.get("name", ""),
        "{{EMAIL}}": profile.get("email", ""),
        "{{PHONE}}": profile.get("phone", ""),
        "{{DATE}}": datetime.now().strftime("%B %d, %Y")
    })

    # Customize greeting (template-based)
    customize_greeting(doc, job)

    # Add keywords without AI
    add_keywords_to_cover_letter_template_based(doc, job.get("keywords", [])[:3])

    # Save to temp file
    temp_docx = utils.create_temp_file(prefix=f"cover_{job_hash}_", suffix=".docx")
    doc.save(temp_docx)

    # Save customized document and convert to PDF
    output_path = utils.save_customized_document(profile, job_hash, temp_docx, is_resume=False)

    return output_path


def customize_cover_letter_basic_replacement(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Basic placeholder replacement cover letter customization (Method 3).

    Args:
        job: Job dictionary
        profile: User profile dictionary
        job_hash: Hash of the job for document naming

    Returns:
        Path to the customized cover letter
    """
    console.print("[cyan]Using basic replacement cover letter customization[/cyan]")

    # Load template cover letter
    profile_dir = Path(profile["profile_dir"])
    cover_letter_docx = profile_dir / profile["cover_letter_docx"]

    if not cover_letter_docx.exists():
        raise FileNotFoundError(f"Cover letter template not found: {cover_letter_docx}")

    # Load the document
    doc = Document(cover_letter_docx)

    # Basic placeholder replacement only
    replacements = {
        "{{JOB_TITLE}}": job.get("title", ""),
        "{{COMPANY}}": job.get("company", ""),
        "{{LOCATION}}": job.get("location", ""),
        "{{NAME}}": profile.get("name", ""),
        "{{EMAIL}}": profile.get("email", ""),
        "{{PHONE}}": profile.get("phone", ""),
        "{{DATE}}": datetime.now().strftime("%B %d, %Y")
    }

    replace_placeholders(doc, replacements)

    # Save to temp file
    temp_docx = utils.create_temp_file(prefix=f"cover_{job_hash}_", suffix=".docx")
    doc.save(temp_docx)

    # Save customized document and convert to PDF
    output_path = utils.save_customized_document(profile, job_hash, temp_docx, is_resume=False)

    return output_path


def create_emergency_text_cover_letter(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Emergency text-based cover letter creation (Method 4).

    Args:
        job: Job dictionary
        profile: User profile dictionary
        job_hash: Hash of the job for document naming

    Returns:
        Path to the emergency cover letter
    """
    console.print("[red]Creating emergency text cover letter[/red]")

    # Create a basic text cover letter
    cover_letter_content = f"""
{profile.get('name', 'Your Name')}
{profile.get('email', 'your.email@example.com')} | {profile.get('phone', 'Your Phone')}
{profile.get('location', 'Your Location')}

{datetime.now().strftime("%B %d, %Y")}

Dear {job.get('company', 'Hiring Manager')},

I am writing to express my strong interest in the {job.get('title', 'position')} role at {job.get('company', 'your company')}.

With my background in {', '.join(profile.get('skills', ['relevant skills'])[:3])}, I am confident that I would be a valuable addition to your team. I have experience with {', '.join(job.get('keywords', ['relevant technologies'])[:3])}, which aligns well with the requirements for this position.

I am particularly excited about the opportunity to contribute to {job.get('company', 'your organization')} and would welcome the chance to discuss how my skills and experience can benefit your team.

Thank you for considering my application. I look forward to hearing from you.

Sincerely,
{profile.get('name', 'Your Name')}

This cover letter was generated automatically for the {job.get('title', 'position')} at {job.get('company', 'the company')}.
"""

    # Save as text file
    temp_txt = utils.create_temp_file(prefix=f"cover_{job_hash}_", suffix=".txt")
    with open(temp_txt, 'w', encoding='utf-8') as f:
        f.write(cover_letter_content)

    console.print(f"[yellow]Emergency text cover letter created: {temp_txt}[/yellow]")
    return temp_txt


def add_keywords_to_cover_letter_template_based(doc: Document, keywords: List[str]) -> None:
    """
    Add keywords to cover letter using template-based approach (no AI).

    Args:
        doc: Document object
        keywords: List of keywords to add
    """
    if not keywords:
        return

    # Find the first body paragraph and add keywords
    for para in doc.paragraphs:
        text = para.text.strip()
        if len(text) > 50 and not text.lower().startswith("dear") and not text.lower().startswith("sincerely"):
            # Add keywords to this paragraph
            para.text = f"{para.text.rstrip('.')}. I have experience with {', '.join(keywords[:2])}."
            break


def customize_resume(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Customize resume for a specific job.
    
    Args:
        job: Job dictionary with title, company, summary, and keywords
        profile: User profile dictionary
        job_hash: Hash of the job for document naming
        
    Returns:
        Path to the customized resume
    """
    console.print("[green]Customizing resume...[/green]")
    
    # Load template resume
    profile_dir = Path(profile["profile_dir"])
    resume_docx = profile_dir / profile["resume_docx"]
    
    if not resume_docx.exists():
        raise FileNotFoundError(f"Resume template not found: {resume_docx}")
    
    # Load the document
    doc = Document(resume_docx)
    
    # Get missing keywords
    profile_skills = set(k.lower() for k in profile.get("skills", []))
    job_keywords = set(k.lower() for k in job.get("keywords", []))
    
    missing_keywords = job_keywords - profile_skills
    console.print(f"[yellow]Missing keywords:[/yellow] {', '.join(missing_keywords)}")
    
    # Only process if there are missing keywords
    if missing_keywords:
        # Limit to top 4 missing keywords
        top_missing = list(missing_keywords)[:4]
        
        # Get all paragraphs with bullet points
        bullet_paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip().startswith('•') or para.text.strip().startswith('-'):
                bullet_paragraphs.append(para)
        
        # Process each missing keyword
        keywords_added = set()
        
        for keyword in top_missing:
            # Try to add the keyword naturally
            if add_keyword_to_bullet(doc, keyword, bullet_paragraphs, profile):
                keywords_added.add(keyword)
        
        # If any keywords are still missing, add them to a skills section
        still_missing = set(top_missing) - keywords_added
        if still_missing:
            add_skills_section(doc, still_missing)
    
    # Save to temp file
    temp_docx = utils.create_temp_file(prefix=f"resume_{job_hash}_", suffix=".docx")
    doc.save(temp_docx)
    
    # Save customized document and convert to PDF
    output_path = utils.save_customized_document(profile, job_hash, temp_docx, is_resume=True)
    
    return output_path


def customize_cover_letter(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Customize cover letter for a specific job.
    
    Args:
        job: Job dictionary with title, company, summary, and keywords
        profile: User profile dictionary
        job_hash: Hash of the job for document naming
        
    Returns:
        Path to the customized cover letter
    """
    console.print("[green]Customizing cover letter...[/green]")
    
    # Load template cover letter
    profile_dir = Path(profile["profile_dir"])
    cover_letter_docx = profile_dir / profile["cover_letter_docx"]
    
    if not cover_letter_docx.exists():
        raise FileNotFoundError(f"Cover letter template not found: {cover_letter_docx}")
    
    # Load the document
    doc = Document(cover_letter_docx)
    
    # Replace placeholders
    replace_placeholders(doc, {
        "{{JOB_TITLE}}": job.get("title", ""),
        "{{COMPANY}}": job.get("company", ""),
        "{{LOCATION}}": job.get("location", ""),
        "{{NAME}}": profile.get("name", ""),
        "{{EMAIL}}": profile.get("email", ""),
        "{{PHONE}}": profile.get("phone", ""),
        "{{DATE}}": datetime.now().strftime("%B %d, %Y")
    })
    
    # Customize greeting
    customize_greeting(doc, job)
    
    # Enhance cover letter with job keywords
    enhance_cover_letter_with_keywords(doc, job, profile)
    
    # Save to temp file
    temp_docx = utils.create_temp_file(prefix=f"cover_{job_hash}_", suffix=".docx")
    doc.save(temp_docx)
    
    # Save customized document and convert to PDF
    output_path = utils.save_customized_document(profile, job_hash, temp_docx, is_resume=False)
    
    return output_path


def add_keyword_to_bullet(doc: Document, keyword: str, bullet_paragraphs: List, profile: Dict) -> bool:
    """
    Add a keyword to an appropriate bullet point.
    
    Args:
        doc: Document object
        keyword: Keyword to add
        bullet_paragraphs: List of paragraphs with bullet points
        
    Returns:
        True if keyword was added, False otherwise
    """
    if not bullet_paragraphs:
        return False
    
    # Try to find the most relevant bullet point
    best_bullet = None
    best_score = -1
    
    for para in bullet_paragraphs:
        # Calculate relevance score (simple word overlap)
        para_text = para.text.lower()
        words = set(re.findall(r'\b\w+\b', para_text))
        
        # Skip if keyword is already in the bullet
        if keyword.lower() in para_text:
            return True
        
        # Calculate relevance based on context
        score = len(words) / 10  # Prefer shorter bullets
        
        if score > best_score:
            best_score = score
            best_bullet = para
    
    # If we found a bullet, try to enhance it with Ollama
    if best_bullet:
        if OLLAMA_AVAILABLE:
            try:
                # Prepare the prompt
                prompt = (
                    f"Here is a resume bullet: \"{best_bullet.text}\". "
                    f"Revise it to smoothly include the keyword \"{keyword}\" if relevant. "
                    f"Return only the revised bullet or 'NONE' if not relevant. "
                    f"Keep the original style and length similar."
                )

                # Generate enhanced bullet
                response = OLLAMA_MODULE.generate(
                    model=profile.get("ollama_model", "mistral"),
                    prompt=prompt,
                    options={
                        "temperature": 0.2,
                        "num_predict": 150
                    }
                )

                enhanced_bullet = response["response"].strip()

                # Check if enhancement was successful
                if enhanced_bullet and enhanced_bullet != "NONE":
                    # Replace the bullet text
                    best_bullet.text = enhanced_bullet
                    console.print(f"[green]Added keyword '{keyword}' to bullet[/green]")
                    return True

            except Exception as e:
                console.print(f"[yellow]Error enhancing bullet with Ollama: {e}[/yellow]")
        else:
            # Fallback: simple keyword insertion
            if keyword.lower() not in best_bullet.text.lower():
                # Add keyword to the end of the bullet
                best_bullet.text = f"{best_bullet.text.rstrip('.')}. Experience with {keyword}."
                console.print(f"[green]Added keyword '{keyword}' to bullet (fallback)[/green]")
                return True
    
    return False


def add_skills_section(doc: Document, keywords: set) -> None:
    """
    Add a skills section with missing keywords.
    
    Args:
        doc: Document object
        keywords: Set of keywords to add
    """
    if not keywords:
        return
    
    # Find "Skills" section or create one at the end
    skills_para = None
    
    for para in doc.paragraphs:
        if para.text.strip().lower() == "skills" or para.text.strip().lower() == "skills:":
            skills_para = para
            break
    
    # If no skills section found, add one
    if not skills_para:
        skills_para = doc.add_paragraph()
        skills_para.text = "Skills"
        # Make it bold
        for run in skills_para.runs:
            run.bold = True
    
    # Add new paragraph for skills
    skills_list = doc.add_paragraph()
    skills_list.text = "• " + ", ".join(keywords)
    
    # Make the text light gray
    for run in skills_list.runs:
        run.font.color.rgb = RGBColor(128, 128, 128)


def replace_placeholders(doc: Document, replacements: Dict[str, str]) -> None:
    """
    Replace placeholders in a document.
    
    Args:
        doc: Document object
        replacements: Dictionary of {placeholder: replacement}
    """
    for para in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, replacement)


def customize_greeting(doc: Document, job: Dict) -> bool:
    """
    Customize the greeting in a cover letter.
    
    Args:
        doc: Document object
        job: Job dictionary
        
    Returns:
        True if greeting was customized, False otherwise
    """
    # Find the greeting paragraph
    greeting_para = None
    
    for para in doc.paragraphs:
        text = para.text.strip().lower()
        if (text.startswith("dear") or 
            text.startswith("to whom") or 
            text.startswith("hello") or 
            text.startswith("hi")):
            greeting_para = para
            break
    
    if not greeting_para:
        return False
    
    # Check if greeting is generic
    generic_greetings = [
        "to whom it may concern",
        "dear hiring manager",
        "dear recruiter",
        "dear sir or madam",
        "dear sir/madam"
    ]
    
    if any(greeting_para.text.strip().lower().startswith(g) for g in generic_greetings):
        # Try to customize greeting with company name
        company = job.get("company", "").strip()
        if company:
            greeting_para.text = f"Dear {company} Hiring Team,"
            return True
    
    return False


def enhance_cover_letter_with_keywords(doc: Document, job: Dict, profile: Dict) -> None:
    """
    Enhance cover letter with job keywords.
    
    Args:
        doc: Document object
        job: Job dictionary
        profile: Profile dictionary
    """
    # Get job keywords
    job_keywords = job.get("keywords", [])
    
    if not job_keywords:
        return
    
    # Find the main body paragraphs (skip header and signature)
    body_paragraphs = []
    in_body = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Skip empty paragraphs
        if not text:
            continue
        
        # Detect start of body (after greeting)
        if not in_body and (text.lower().startswith("dear") or 
                           text.lower().startswith("to whom") or 
                           text.lower().startswith("hello")):
            in_body = True
            continue
        
        # Detect end of body (signature)
        if in_body and (text.lower().startswith("sincerely") or 
                       text.lower().startswith("regards") or 
                       text.lower().startswith("thank you") or 
                       text.lower() == profile.get("name", "").lower()):
            in_body = False
        
        # Add paragraph to body if in body section
        if in_body:
            body_paragraphs.append(para)
    
    # If we found body paragraphs, try to enhance one with keywords
    if body_paragraphs:
        if OLLAMA_AVAILABLE:
            try:
                # Choose the longest paragraph to enhance
                target_para = max(body_paragraphs, key=lambda p: len(p.text))

                # Prepare the prompt
                prompt = (
                    f"Here is a paragraph from a cover letter: \"{target_para.text}\". "
                    f"Revise it to naturally include 2-3 of these keywords: {', '.join(job_keywords[:5])}. "
                    f"Keep the same tone and length. Return only the revised paragraph."
                )

                # Generate enhanced paragraph
                response = OLLAMA_MODULE.generate(
                    model=profile.get("ollama_model", "mistral"),
                    prompt=prompt,
                    options={
                        "temperature": 0.3,
                        "num_predict": 300
                    }
                )

                enhanced_para = response["response"].strip()

                # Check if enhancement was successful
                if enhanced_para and len(enhanced_para) > 20:
                    # Replace the paragraph text
                    target_para.text = enhanced_para
                    console.print(f"[green]Enhanced cover letter with job keywords[/green]")

            except Exception as e:
                console.print(f"[yellow]Error enhancing cover letter with Ollama: {e}[/yellow]")
        else:
            # Fallback: add keywords to the end of the first paragraph
            if body_paragraphs:
                target_para = body_paragraphs[0]
                keywords_to_add = job_keywords[:3]  # Add up to 3 keywords
                if keywords_to_add:
                    # Add keywords naturally to the paragraph
                    target_para.text = f"{target_para.text.rstrip('.')}. I have experience with {', '.join(keywords_to_add)}."
                    console.print(f"[green]Enhanced cover letter with job keywords (fallback)[/green]")


def flatten_document_styles(doc: Document) -> None:
    """Flatten document styles to ensure consistent formatting."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Reset any complex formatting
            run.font.name = 'Calibri'
            run.font.size = None
            run.font.bold = None
            run.font.italic = None
            run.font.underline = None

class DocumentGenerator:
    """
    Document generator class for creating customized resumes and cover letters.
    Provides a unified interface for document customization.
    """
    
    def __init__(self, profile_name: str = "default"):
        self.profile_name = profile_name
        self.console = Console()
    
    def generate_cover_letter(self, job: Dict, profile: Dict) -> str:
        """Generate a customized cover letter for a job."""
        try:
            job_hash = utils.hash_job(job)
            return customize_cover_letter(job, profile, job_hash)
        except Exception as e:
            self.console.print(f"[red]Error generating cover letter: {e}[/red]")
            raise
    
    def generate_resume(self, job: Dict, profile: Dict) -> str:
        """Generate a customized resume for a job."""
        try:
            job_hash = utils.hash_job(job)
            return customize_resume(job, profile, job_hash)
        except Exception as e:
            self.console.print(f"[red]Error generating resume: {e}[/red]")
            raise
    
    def get_available_templates(self) -> List[str]:
        """Get list of available document templates."""
        return ["default", "professional", "creative", "minimal"]
    
    def format_document(self, document_path: str, format_type: str = "docx") -> str:
        """Format a document to the specified format."""
        # For now, just return the original path
        return document_path
    
    def use_custom_template(self, template_path: str) -> bool:
        """Use a custom template for document generation."""
        # For now, just return True
        return True

if __name__ == "__main__":
    # Test code
    test_job = {
        "title": "Data Analyst",
        "company": "Test Company",
        "location": "Toronto, ON",
        "summary": "Looking for a data analyst with Python, SQL, and Power BI experience.",
        "url": "https://example.com/job",
        "keywords": ["Python", "SQL", "Power BI", "Excel", "Data Visualization", "Statistical Analysis"]
    }
    
    test_profile = {
        "name": "Nirajan Khadka",
        "location": "Mississauga, ON",
        "email": "Nirajan.Tech@gmail.com",
        "phone": "437-344-5361",
        "profile_name": "Nirajan",
        "profile_dir": "profiles/Nirajan",
        "resume_docx": "Nirajan_Resume.docx",
        "cover_letter_docx": "Nirajan_CoverLetter.docx",
        "resume_pdf": "Nirajan_Resume.pdf",
        "cover_letter_pdf": "Nirajan_CoverLetter.pdf",
        "skills": ["Python", "Pandas", "Excel", "Statistical Analysis"],
        "ollama_model": "mistral:7b"
    }
    
    resume_path, cover_letter_path = customize(test_job, test_profile)
    print(f"Customized resume: {resume_path}")
    print(f"Customized cover letter: {cover_letter_path}")
