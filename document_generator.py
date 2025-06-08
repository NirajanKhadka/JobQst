import json
import os
import re
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

import ollama
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from rich.console import Console

import utils

# Initialize console for rich output
console = Console()

def customize(job: Dict, profile: Dict) -> Tuple[str, str]:
    """
    Customize resume and cover letter for a specific job.
    
    Args:
        job: Job dictionary with title, company, summary, and keywords
        profile: User profile dictionary
        
    Returns:
        Tuple of (resume_path, cover_letter_path)
    """
    console.print(f"[bold]Customizing documents for:[/bold] {job['title']} at {job['company']}")
    
    # Get job hash for document naming
    job_hash = utils.hash_job(job)
    
    # Customize resume
    resume_path = customize_resume(job, profile, job_hash)
    
    # Customize cover letter
    cover_letter_path = customize_cover_letter(job, profile, job_hash)
    
    return resume_path, cover_letter_path


def customize_resume(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Customize resume for a specific job.
    
    Args:
        job: Job dictionary with title, company, summary, and keywords
        profile: User profile dictionary
        job_hash: Hash of the job for document naming
        
    Returns:
        Path to the customized resume
    """
    console.print("[green]Customizing resume...[/green]")
    
    # Load template resume
    profile_dir = Path(profile["profile_dir"])
    resume_docx = profile_dir / profile["resume_docx"]
    
    if not resume_docx.exists():
        raise FileNotFoundError(f"Resume template not found: {resume_docx}")
    
    # Load the document
    doc = Document(resume_docx)
    
    # Get missing keywords
    profile_skills = set(k.lower() for k in profile.get("skills", []))
    job_keywords = set(k.lower() for k in job.get("keywords", []))
    
    missing_keywords = job_keywords - profile_skills
    console.print(f"[yellow]Missing keywords:[/yellow] {', '.join(missing_keywords)}")
    
    # Only process if there are missing keywords
    if missing_keywords:
        # Limit to top 4 missing keywords
        top_missing = list(missing_keywords)[:4]
        
        # Get all paragraphs with bullet points
        bullet_paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip().startswith('•') or para.text.strip().startswith('-'):
                bullet_paragraphs.append(para)
        
        # Process each missing keyword
        keywords_added = set()
        
        for keyword in top_missing:
            # Try to add the keyword naturally
            if add_keyword_to_bullet(doc, keyword, bullet_paragraphs):
                keywords_added.add(keyword)
        
        # If any keywords are still missing, add them to a skills section
        still_missing = set(top_missing) - keywords_added
        if still_missing:
            add_skills_section(doc, still_missing)
    
    # Save to temp file
    temp_docx = utils.create_temp_file(prefix=f"resume_{job_hash}_", suffix=".docx")
    doc.save(temp_docx)
    
    # Save customized document and convert to PDF
    output_path = utils.save_customized_document(profile, job_hash, temp_docx, is_resume=True)
    
    return output_path


def customize_cover_letter(job: Dict, profile: Dict, job_hash: str) -> str:
    """
    Customize cover letter for a specific job.
    
    Args:
        job: Job dictionary with title, company, summary, and keywords
        profile: User profile dictionary
        job_hash: Hash of the job for document naming
        
    Returns:
        Path to the customized cover letter
    """
    console.print("[green]Customizing cover letter...[/green]")
    
    # Load template cover letter
    profile_dir = Path(profile["profile_dir"])
    cover_letter_docx = profile_dir / profile["cover_letter_docx"]
    
    if not cover_letter_docx.exists():
        raise FileNotFoundError(f"Cover letter template not found: {cover_letter_docx}")
    
    # Load the document
    doc = Document(cover_letter_docx)
    
    # Replace placeholders
    replace_placeholders(doc, {
        "{{JOB_TITLE}}": job.get("title", ""),
        "{{COMPANY}}": job.get("company", ""),
        "{{LOCATION}}": job.get("location", ""),
        "{{NAME}}": profile.get("name", ""),
        "{{EMAIL}}": profile.get("email", ""),
        "{{PHONE}}": profile.get("phone", ""),
        "{{DATE}}": datetime.now().strftime("%B %d, %Y")
    })
    
    # Customize greeting
    customize_greeting(doc, job)
    
    # Enhance cover letter with job keywords
    enhance_cover_letter_with_keywords(doc, job, profile)
    
    # Save to temp file
    temp_docx = utils.create_temp_file(prefix=f"cover_{job_hash}_", suffix=".docx")
    doc.save(temp_docx)
    
    # Save customized document and convert to PDF
    output_path = utils.save_customized_document(profile, job_hash, temp_docx, is_resume=False)
    
    return output_path


def add_keyword_to_bullet(doc: Document, keyword: str, bullet_paragraphs: List) -> bool:
    """
    Add a keyword to an appropriate bullet point.
    
    Args:
        doc: Document object
        keyword: Keyword to add
        bullet_paragraphs: List of paragraphs with bullet points
        
    Returns:
        True if keyword was added, False otherwise
    """
    if not bullet_paragraphs:
        return False
    
    # Try to find the most relevant bullet point
    best_bullet = None
    best_score = -1
    
    for para in bullet_paragraphs:
        # Calculate relevance score (simple word overlap)
        para_text = para.text.lower()
        words = set(re.findall(r'\b\w+\b', para_text))
        
        # Skip if keyword is already in the bullet
        if keyword.lower() in para_text:
            return True
        
        # Calculate relevance based on context
        score = len(words) / 10  # Prefer shorter bullets
        
        if score > best_score:
            best_score = score
            best_bullet = para
    
    # If we found a bullet, try to enhance it with Ollama
    if best_bullet:
        try:
            # Prepare the prompt
            prompt = (
                f"Here is a resume bullet: \"{best_bullet.text}\". "
                f"Revise it to smoothly include the keyword \"{keyword}\" if relevant. "
                f"Return only the revised bullet or 'NONE' if not relevant. "
                f"Keep the original style and length similar."
            )
            
            # Generate enhanced bullet
            response = ollama.generate(
                model=profile.get("ollama_model", "mistral:7b"),
                prompt=prompt,
                temperature=0.2,
                max_tokens=150
            )
            
            enhanced_bullet = response["response"].strip()
            
            # Check if enhancement was successful
            if enhanced_bullet and enhanced_bullet != "NONE":
                # Replace the bullet text
                best_bullet.text = enhanced_bullet
                console.print(f"[green]Added keyword '{keyword}' to bullet[/green]")
                return True
        
        except Exception as e:
            console.print(f"[yellow]Error enhancing bullet with Ollama: {e}[/yellow]")
    
    return False


def add_skills_section(doc: Document, keywords: set) -> None:
    """
    Add a skills section with missing keywords.
    
    Args:
        doc: Document object
        keywords: Set of keywords to add
    """
    if not keywords:
        return
    
    # Find "Skills" section or create one at the end
    skills_para = None
    
    for para in doc.paragraphs:
        if para.text.strip().lower() == "skills" or para.text.strip().lower() == "skills:":
            skills_para = para
            break
    
    # If no skills section found, add one
    if not skills_para:
        skills_para = doc.add_paragraph()
        skills_para.text = "Skills"
        # Make it bold
        for run in skills_para.runs:
            run.bold = True
    
    # Add new paragraph for skills
    skills_list = doc.add_paragraph()
    skills_list.text = "• " + ", ".join(keywords)
    
    # Make the text light gray
    for run in skills_list.runs:
        run.font.color.rgb = RGBColor(128, 128, 128)


def replace_placeholders(doc: Document, replacements: Dict[str, str]) -> None:
    """
    Replace placeholders in a document.
    
    Args:
        doc: Document object
        replacements: Dictionary of {placeholder: replacement}
    """
    for para in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, replacement)


def customize_greeting(doc: Document, job: Dict) -> bool:
    """
    Customize the greeting in a cover letter.
    
    Args:
        doc: Document object
        job: Job dictionary
        
    Returns:
        True if greeting was customized, False otherwise
    """
    # Find the greeting paragraph
    greeting_para = None
    
    for para in doc.paragraphs:
        text = para.text.strip().lower()
        if (text.startswith("dear") or 
            text.startswith("to whom") or 
            text.startswith("hello") or 
            text.startswith("hi")):
            greeting_para = para
            break
    
    if not greeting_para:
        return False
    
    # Check if greeting is generic
    generic_greetings = [
        "to whom it may concern",
        "dear hiring manager",
        "dear recruiter",
        "dear sir or madam",
        "dear sir/madam"
    ]
    
    if any(greeting_para.text.strip().lower().startswith(g) for g in generic_greetings):
        # Try to customize greeting with company name
        company = job.get("company", "").strip()
        if company:
            greeting_para.text = f"Dear {company} Hiring Team,"
            return True
    
    return False


def enhance_cover_letter_with_keywords(doc: Document, job: Dict, profile: Dict) -> None:
    """
    Enhance cover letter with job keywords.
    
    Args:
        doc: Document object
        job: Job dictionary
        profile: Profile dictionary
    """
    # Get job keywords
    job_keywords = job.get("keywords", [])
    
    if not job_keywords:
        return
    
    # Find the main body paragraphs (skip header and signature)
    body_paragraphs = []
    in_body = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Skip empty paragraphs
        if not text:
            continue
        
        # Detect start of body (after greeting)
        if not in_body and (text.lower().startswith("dear") or 
                           text.lower().startswith("to whom") or 
                           text.lower().startswith("hello")):
            in_body = True
            continue
        
        # Detect end of body (signature)
        if in_body and (text.lower().startswith("sincerely") or 
                       text.lower().startswith("regards") or 
                       text.lower().startswith("thank you") or 
                       text.lower() == profile.get("name", "").lower()):
            in_body = False
        
        # Add paragraph to body if in body section
        if in_body:
            body_paragraphs.append(para)
    
    # If we found body paragraphs, try to enhance one with keywords
    if body_paragraphs:
        try:
            # Choose the longest paragraph to enhance
            target_para = max(body_paragraphs, key=lambda p: len(p.text))
            
            # Prepare the prompt
            prompt = (
                f"Here is a paragraph from a cover letter: \"{target_para.text}\". "
                f"Revise it to naturally include 2-3 of these keywords: {', '.join(job_keywords[:5])}. "
                f"Keep the same tone and length. Return only the revised paragraph."
            )
            
            # Generate enhanced paragraph
            response = ollama.generate(
                model=profile.get("ollama_model", "mistral:7b"),
                prompt=prompt,
                temperature=0.3,
                max_tokens=300
            )
            
            enhanced_para = response["response"].strip()
            
            # Check if enhancement was successful
            if enhanced_para and len(enhanced_para) > 20:
                # Replace the paragraph text
                target_para.text = enhanced_para
                console.print(f"[green]Enhanced cover letter with job keywords[/green]")
        
        except Exception as e:
            console.print(f"[yellow]Error enhancing cover letter with Ollama: {e}[/yellow]")


def flatten_document_styles(doc: Document) -> None:
    """
    Flatten styles in a document (convert to plain text).
    
    Args:
        doc: Document object
    """
    # Process all paragraphs
    for para in doc.paragraphs:
        # Skip empty paragraphs
        if not para.text.strip():
            continue
        
        # Get the text
        text = para.text
        
        # Clear the paragraph
        para.clear()
        
        # Add the text back as a single run
        para.add_run(text)


if __name__ == "__main__":
    # Test code
    test_job = {
        "title": "Data Analyst",
        "company": "Test Company",
        "location": "Toronto, ON",
        "summary": "Looking for a data analyst with Python, SQL, and Power BI experience.",
        "url": "https://example.com/job",
        "keywords": ["Python", "SQL", "Power BI", "Excel", "Data Visualization", "Statistical Analysis"]
    }
    
    test_profile = {
        "name": "Nirajan Khadka",
        "location": "Mississauga, ON",
        "email": "Nirajan.Tech@gmail.com",
        "phone": "437-344-5361",
        "profile_name": "Nirajan",
        "profile_dir": "profiles/Nirajan",
        "resume_docx": "Nirajan_Resume.docx",
        "cover_letter_docx": "Nirajan_CoverLetter.docx",
        "resume_pdf": "Nirajan_Resume.pdf",
        "cover_letter_pdf": "Nirajan_CoverLetter.pdf",
        "skills": ["Python", "Pandas", "Excel", "Statistical Analysis"],
        "ollama_model": "mistral:7b"
    }
    
    resume_path, cover_letter_path = customize(test_job, test_profile)
    print(f"Customized resume: {resume_path}")
    print(f"Customized cover letter: {cover_letter_path}")
